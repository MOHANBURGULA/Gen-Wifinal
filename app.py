from flask import Flask, request, send_file
from docx import Document
import io

app = Flask(__name__)

@app.route('/generate-document', methods=['POST'])
def generate_document():
    # Create a new Word document
    document = Document()
    
    # Get form data
    project_name = request.form.get('projectName')
    theme = request.form.get('theme')
    members = request.form.get('members')
    abstract = request.form.get('abstract')
    mentor = request.form.get('mentor')
    title = request.form.get('title')
    author = request.form.get('author')
    content = request.form.get('content')

    # Check if any field is empty (optional)
    if not all([project_name, theme, members, abstract, mentor, title, author, content]):
        return {'error': 'All form fields must be filled in'}, 400

    # Add data to the Word document
    document.add_heading(title, level=1)
    document.add_paragraph(f"Project Name: {project_name}")
    document.add_paragraph(f"Project Theme: {theme}")
    document.add_paragraph(f"Members: {members}")
    document.add_paragraph(f"Abstract: {abstract}")
    document.add_paragraph(f"Mentor: {mentor}")
    document.add_paragraph(f"Incharge: {author}")
    document.add_paragraph(f"Content: {content}")

    # Save document to a BytesIO stream
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # Send the file as a response
    return send_file(file_stream, as_attachment=True, download_name='Generated_Document.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    app.run(debug=True)
